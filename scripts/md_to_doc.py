#!/usr/bin/env python3
"""Convert the WAI market research markdown report to a Word .doc(x) file."""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

src = sys.argv[1]
out = sys.argv[2]

with open(src, encoding="utf-8") as f:
    text = f.read()

# Strip YAML front matter
text = re.sub(r"^---\n.*?\n---\n", "", text, count=1, flags=re.DOTALL)
lines = text.split("\n")

doc = Document()

# Base styling
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def add_inline(paragraph, content):
    """Render bold (**..**), italic (*..*), inline code (`..`) and links."""
    # Tokenize on markdown emphasis / code / links
    pattern = r"(\*\*.+?\*\*|\*.+?\*|`.+?`|\[.+?\]\(.+?\))"
    for tok in re.split(pattern, content):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            run = paragraph.add_run(tok[1:-1])
            run.italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = "Consolas"
        elif tok.startswith("["):
            m = re.match(r"\[(.+?)\]\((.+?)\)", tok)
            if m:
                run = paragraph.add_run(m.group(1) + " (" + m.group(2) + ")")
                run.font.color.rgb = RGBColor(0x1A, 0x5F, 0xB4)
            else:
                paragraph.add_run(tok)
        else:
            paragraph.add_run(tok)


def flush_table(rows):
    if not rows:
        return
    # rows: list of list[str]; first row is header, drop separator rows
    body = [r for r in rows if not re.match(r"^[\s\-:|]+$", "|".join(r))]
    if not body:
        return
    ncol = max(len(r) for r in body)
    table = doc.add_table(rows=0, cols=ncol)
    table.style = "Light Grid Accent 1"
    for i, r in enumerate(body):
        cells = table.add_row().cells
        for j in range(ncol):
            val = r[j] if j < len(r) else ""
            cell = cells[j]
            cell.paragraphs[0].text = ""
            add_inline(cell.paragraphs[0], val.strip())
            if i == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True


table_buf = []
i = 0
while i < len(lines):
    line = lines[i]

    # Table detection
    if line.strip().startswith("|") and line.strip().endswith("|"):
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        table_buf.append(cells)
        i += 1
        continue
    elif table_buf:
        flush_table(table_buf)
        table_buf = []

    stripped = line.strip()

    if not stripped:
        i += 1
        continue
    if stripped == "---":
        i += 1
        continue

    if stripped.startswith("#"):
        m = re.match(r"(#+)\s*(.*)", stripped)
        level = min(len(m.group(1)), 4)
        heading = doc.add_heading(level=level)
        add_inline(heading, m.group(2))
    elif re.match(r"^[-*]\s+", stripped):
        p = doc.add_paragraph(style="List Bullet")
        add_inline(p, re.sub(r"^[-*]\s+", "", stripped))
    elif re.match(r"^\d+\.\s+", stripped):
        p = doc.add_paragraph(style="List Number")
        add_inline(p, re.sub(r"^\d+\.\s+", "", stripped))
    else:
        p = doc.add_paragraph()
        add_inline(p, stripped)
    i += 1

if table_buf:
    flush_table(table_buf)

doc.save(out)
print("Saved", out)
